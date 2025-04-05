import os
import tempfile
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import img2pdf
from PIL import Image
import zipfile
import io
import subprocess
import platform
import shutil
import logging

logging.basicConfig(level=logging.INFO)

class PDFTransformer: # Poderia ser renomeado para PDFProcessor ou algo similar

    def __init__(self):
        self.ocrmypdf_cmd = 'ocrmypdf'
        self.ocrmypdf_installed = self._check_ocrmypdf_installed()
        self.libreoffice_path = self._find_libreoffice()
        # --- Lógica do Compressor adicionada ---
        self.gs_cmd = 'gswin64c' if platform.system() == 'Windows' else 'gs'
        self.gs_available = self._check_gs_available()
        # --- Fim da adição ---

    def _check_ocrmypdf_installed(self):
        # ... (código existente inalterado) ...
        try:
            subprocess.run([self.ocrmypdf_cmd, '--version'],
                           stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                           check=True, text=True, timeout=5)
            logging.info("OCRmyPDF detectado.")
            return True
        except (subprocess.SubprocessError, FileNotFoundError, Exception) as e:
            logging.warning(f"OCRmyPDF não detectado ou erro ao verificar: {e}")
            return False

    def _apply_ocrmypdf(self, input_path, output_path, language='por'):
         # ... (código existente inalterado, mas com melhor log de erro) ...
         if not self.ocrmypdf_installed:
             logging.error("Tentativa de usar OCRmyPDF, mas não está instalado.")
             return False
         # Usa diretório temporário para evitar conflitos de nome
         with tempfile.TemporaryDirectory() as temp_ocr_dir:
             temp_input = os.path.join(temp_ocr_dir, "input_ocr.pdf")
             shutil.copy2(input_path, temp_input)
             # Define o caminho de saída dentro do mesmo diretório temporário
             temp_output = os.path.join(temp_ocr_dir, "output_ocr.pdf")

             args = [
                 self.ocrmypdf_cmd,
                 '--force-ocr',
                 '--optimize', '1',
                 '--output-type', 'pdf',
                 '--jobs', '2',
                 '-l', language,
                 temp_input,
                 temp_output # Saída temporária
             ]
             try:
                 process = subprocess.run(args, capture_output=True, text=True, check=False, timeout=300) # Timeout 5 min

                 if process.returncode != 0:
                     logging.error(f"Erro no OCRmyPDF (código {process.returncode}):")
                     logging.error(f"Stderr: {process.stderr}")
                     logging.error(f"Stdout: {process.stdout}")
                     return False

                 # Copia o resultado bem-sucedido para o caminho de saída final
                 if os.path.exists(temp_output):
                     shutil.move(temp_output, output_path)
                     logging.info(f"OCR aplicado com sucesso e salvo em {output_path}")
                     return True
                 else:
                     logging.error("OCRmyPDF terminou sem erros, mas o arquivo de saída não foi encontrado.")
                     return False

             except subprocess.TimeoutExpired:
                  logging.error("OCRmyPDF excedeu o tempo limite.")
                  return False
             except Exception as e:
                  logging.error(f"Erro ao executar OCRmyPDF: {str(e)}")
                  return False
         # temp_ocr_dir é limpo automaticamente


    def _find_libreoffice(self):
        # ... (código existente inalterado) ...
        for cmd in ['soffice', 'libreoffice']:
            try:
                subprocess.run([cmd, '--version'], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=5)
                logging.info(f"LibreOffice detectado usando o comando: '{cmd}'")
                return cmd
            except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
                continue
        logging.warning("Comando do LibreOffice ('soffice' ou 'libreoffice') não encontrado no PATH.")
        return None

    # --- Métodos do Compressor Adicionados/Adaptados ---
    def _check_gs_available(self):
        """Verifica se o Ghostscript está acessível."""
        try:
            subprocess.run([self.gs_cmd, '--version'],
                           stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                           check=True, text=True, timeout=5)
            logging.info(f"Ghostscript detectado usando o comando: '{self.gs_cmd}'")
            return True
        except (subprocess.SubprocessError, FileNotFoundError, Exception) as e:
            logging.warning(f"Ghostscript ({self.gs_cmd}) não detectado ou erro: {e}")
            return False

    def _compress_pdf_gs(self, input_file_path, output_file_path, power=3):
        """Comprime o PDF usando Ghostscript (método interno)."""
        if not self.gs_available:
            logging.error("Ghostscript não está disponível. Não é possível comprimir.")
            return False
        quality = {
            0: '/default',  # Qualidade próxima ao original, menor compressão
            1: '/prepress', # Alta qualidade (300 dpi)
            2: '/printer',  # Qualidade de impressão (300 dpi)
            3: '/ebook',    # Qualidade média (150 dpi)
            4: '/screen'    # Baixa qualidade, alta compressão (72 dpi)
        }
        # Garante que power esteja no range válido
        power = max(0, min(4, power))

        command = [
            self.gs_cmd, '-sDEVICE=pdfwrite', '-dCompatibilityLevel=1.4',
            f'-dPDFSETTINGS={quality[power]}',
            '-dNOPAUSE', '-dQUIET', '-dBATCH',
            f'-sOutputFile={output_file_path}',
            input_file_path
        ]
        logging.info(f"Executando compressão Ghostscript: {' '.join(command)}")
        try:
            process = subprocess.run(command, capture_output=True, check=False, text=True, timeout=300) # Timeout 5 min
            if process.returncode != 0:
                 logging.error(f"Erro no Ghostscript (código {process.returncode}):")
                 logging.error(f"Stderr: {process.stderr}")
                 logging.error(f"Stdout: {process.stdout}")
                 # Tenta remover arquivo de saída potencialmente corrompido
                 if os.path.exists(output_file_path): os.unlink(output_file_path)
                 return False

            # Verifica se o arquivo foi criado e tem algum tamanho
            if os.path.exists(output_file_path) and os.path.getsize(output_file_path) > 0:
                logging.info("Compressão com Ghostscript concluída com sucesso.")
                return True
            else:
                logging.error("Ghostscript finalizou sem erro, mas o arquivo de saída não foi criado ou está vazio.")
                return False
        except subprocess.TimeoutExpired:
             logging.error("Compressão com Ghostscript excedeu o tempo limite.")
             if os.path.exists(output_file_path): os.unlink(output_file_path) # Limpeza
             return False
        except Exception as e:
            logging.error(f"Erro inesperado ao executar Ghostscript: {str(e)}")
            if os.path.exists(output_file_path): os.unlink(output_file_path) # Limpeza
            return False

    def process_compression_ocr(self, file_bytes, compression_level=3, apply_ocr=False, ocr_language='por'):
        """
        Processa um PDF: comprime e/ou aplica OCR.

        :param file_bytes: Bytes do arquivo PDF de entrada.
        :param compression_level: Nível de compressão Ghostscript (0-4). Ignorado se apply_ocr=True e compressão não solicitada explicitamente.
        :param apply_ocr: Boolean, se True aplica OCR.
        :param ocr_language: Idioma para OCR.
        :return: Tupla (success, processed_bytes, original_size_mb, final_size_mb)
        """
        original_size_mb = len(file_bytes) / 1024 / 1024
        processed_bytes = None
        final_size_mb = 0
        success = False

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, "input.pdf")
            with open(input_path, "wb") as f:
                f.write(file_bytes)

            current_step_output = input_path # Começa com o arquivo original
            step_successful = True

            # Etapa 1: Compressão (se compression_level > -1, por exemplo)
            # Permitir compress_level 0 (qualidade default)
            if compression_level >= 0:
                logging.info(f"Iniciando Etapa de Compressão (Nível {compression_level})...")
                compressed_path = os.path.join(temp_dir, "compressed.pdf")
                step_successful = self._compress_pdf_gs(current_step_output, compressed_path, compression_level)
                if step_successful:
                    current_step_output = compressed_path # Atualiza o arquivo a ser processado
                    logging.info("Compressão bem-sucedida.")
                else:
                    logging.error("Falha na etapa de compressão.")
                    # Decide se continua sem compressão ou falha tudo? Vamos falhar tudo por enquanto.
                    return False, None, original_size_mb, 0


            # Etapa 2: OCR (se solicitado e etapa anterior ok)
            if step_successful and apply_ocr:
                logging.info("Iniciando Etapa de OCR...")
                ocr_output_path = os.path.join(temp_dir, "ocr_output.pdf")
                step_successful = self._apply_ocrmypdf(current_step_output, ocr_output_path, ocr_language)
                if step_successful:
                    current_step_output = ocr_output_path # Atualiza o arquivo final
                    logging.info("OCR bem-sucedido.")
                else:
                    logging.warning("Falha na etapa de OCR. O resultado pode não conter OCR.")
                    # Decide se continua com o arquivo anterior ou falha?
                    # Vamos continuar com o resultado da etapa anterior (compressão, se houve)
                    step_successful = True # Marca como sucesso parcial (sem OCR)


            # Leitura do resultado final
            if step_successful and os.path.exists(current_step_output):
                 with open(current_step_output, 'rb') as f:
                     processed_bytes = f.read()
                 final_size_mb = len(processed_bytes) / 1024 / 1024
                 success = True
            else:
                 logging.error("Falha no processamento ou arquivo final não encontrado.")


        return success, processed_bytes, original_size_mb, final_size_mb
    # --- Fim dos métodos do Compressor ---


    # --- Métodos de Transformação (existentes) ---
    def image_to_pdf(self, image_files_bytes, output_pdf_path):
        # ... (código existente inalterado) ...
        try:
            valid_image_bytes = []
            # ... (validação de imagem) ...
            for img_bytes in image_files_bytes:
                 try:
                      img = Image.open(io.BytesIO(img_bytes))
                      img.verify()
                      img = Image.open(io.BytesIO(img_bytes))
                      if img.mode in ('P', 'RGBA', 'LA'):
                           try: img = img.convert('RGB')
                           except Exception as convert_err:
                                logging.warning(f"Não foi possível converter imagem para RGB, pulando: {convert_err}")
                                continue
                      byte_io = io.BytesIO()
                      save_format = 'JPEG' if img.format != 'PNG' else 'PNG'
                      img.save(byte_io, format=save_format)
                      valid_image_bytes.append(byte_io.getvalue())
                      img.close()
                 except Exception as img_err:
                      logging.warning(f"Ignorando arquivo de imagem inválido ou não suportado: {img_err}")
                      continue

            if not valid_image_bytes:
                 logging.error("Nenhuma imagem válida fornecida para conversão.")
                 return False
            with open(output_pdf_path, "wb") as f:
                f.write(img2pdf.convert(valid_image_bytes))
            logging.info(f"Imagens convertidas para PDF em: {output_pdf_path}")
            return True
        except Exception as e:
            logging.error(f"Erro ao converter imagem para PDF: {str(e)}")
            return False

    def pdf_to_docx(self, input_pdf_path, output_docx_path, apply_ocr=False):
        # ... (código existente inalterado, mas usa self._apply_ocrmypdf) ...
        pdf_to_process = input_pdf_path
        temp_ocr_pdf = None
        ocr_applied_successfully = False

        try:
            # Cria diretório temporário para arquivos intermediários desta operação
            with tempfile.TemporaryDirectory() as temp_docx_dir:
                 if apply_ocr:
                      if not self.ocrmypdf_installed:
                           logging.warning("OCR solicitado, mas OCRmyPDF não está instalado.")
                      else:
                           logging.info("Aplicando OCR antes da conversão para DOCX...")
                           temp_ocr_pdf_path = os.path.join(temp_docx_dir, "ocr_temp.pdf")
                           # Chama o método _apply_ocrmypdf da própria classe
                           if self._apply_ocrmypdf(input_pdf_path, temp_ocr_pdf_path):
                                pdf_to_process = temp_ocr_pdf_path
                                ocr_applied_successfully = True
                                logging.info("OCR aplicado. Prosseguindo com a conversão.")
                           else:
                                logging.warning("Falha ao aplicar OCR. Prosseguindo com o PDF original.")

                 logging.info(f"Iniciando conversão de {os.path.basename(pdf_to_process)} para DOCX.")
                 doc = fitz.open(pdf_to_process)
                 document = Document()
                 has_content = False

                 # ... (lógica de extração de texto e imagem) ...
                 for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text("text")
                    if text.strip():
                         document.add_paragraph(text)
                         has_content = True
                    image_list = page.get_images(full=True)
                    for img_index, img_info in enumerate(image_list):
                         xref = img_info[0]
                         try:
                              base_image = doc.extract_image(xref)
                              image_bytes = base_image["image"]
                              image_ext = base_image["ext"]
                              # Salva imagem temporariamente DENTRO do temp_docx_dir
                              temp_img_path = os.path.join(temp_docx_dir, f"page{page_num}_img{img_index}.{image_ext}")
                              with open(temp_img_path, "wb") as temp_img:
                                   temp_img.write(image_bytes)
                              # Adiciona imagem
                              try:
                                   document.add_picture(temp_img_path, width=Inches(6.0))
                                   has_content = True
                              except Exception as pic_err:
                                   logging.warning(f"Não foi possível adicionar a imagem {img_index} da pág {page_num+1}: {pic_err}")
                              # Limpeza da imagem individual não é mais necessária aqui, o diretório será limpo
                         except Exception as img_extract_err:
                              logging.warning(f"Erro ao extrair imagem {img_index} da pág {page_num+1}: {img_extract_err}")

                    if page_num < len(doc) - 1: document.add_page_break()

                 doc.close()

                 if not has_content:
                      if apply_ocr and not ocr_applied_successfully:
                           logging.warning("Nenhum texto/imagem extraído e OCR falhou. DOCX pode estar vazio.")
                      else:
                           logging.warning("Nenhum texto/imagem extraído. DOCX pode estar vazio.")

                 document.save(output_docx_path)
                 logging.info(f"PDF convertido para DOCX em: {output_docx_path}")
                 return True

        except Exception as e:
            logging.error(f"Erro ao converter PDF para DOCX: {str(e)}")
            return False
        # temp_docx_dir é limpo automaticamente


    def pdf_to_image(self, input_pdf_path, output_folder, image_format='png'):
        # ... (código existente inalterado, talvez aumentar zoom padrão) ...
        generated_images = []
        try:
            if not os.path.exists(output_folder):
                 os.makedirs(output_folder)
            doc = fitz.open(input_pdf_path)
            logging.info(f"Convertendo {len(doc)} páginas de PDF para {image_format.upper()}...")
            for page_num in range(len(doc)):
                 page = doc.load_page(page_num)
                 zoom = 2.0 # Zoom de 2x (aprox 144 DPI) - ajuste conforme necessário
                 mat = fitz.Matrix(zoom, zoom)
                 pix = page.get_pixmap(matrix=mat, alpha=False) # alpha=False pode ajudar com PNGs
                 output_image_path = os.path.join(output_folder, f"pagina_{page_num + 1}.{image_format}")
                 pix.save(output_image_path)
                 generated_images.append(output_image_path)
            doc.close()
            logging.info(f"{len(generated_images)} imagens geradas em: {output_folder}")
            return generated_images
        except Exception as e:
            logging.error(f"Erro ao converter PDF para imagem: {str(e)}")
            return None

    def create_zip_from_files(self, file_paths, output_zip_path):
        # ... (código existente inalterado) ...
        try:
            with zipfile.ZipFile(output_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                 for file_path in file_paths:
                     zipf.write(file_path, os.path.basename(file_path))
            logging.info(f"Arquivo ZIP criado com sucesso em: {output_zip_path}")
            return True
        except Exception as e:
            logging.error(f"Erro ao criar arquivo ZIP: {str(e)}")
            return False

    def document_to_pdf(self, input_doc_path, output_pdf_path):
        # ... (código existente inalterado) ...
        if not self.libreoffice_path:
            logging.error("LibreOffice não encontrado.")
            return False
        output_dir = os.path.dirname(output_pdf_path)
        input_basename = os.path.splitext(os.path.basename(input_doc_path))[0]
        expected_libreoffice_output = os.path.join(output_dir, f"{input_basename}.pdf")
        # ... (limpeza de arquivos antigos) ...
        if os.path.exists(expected_libreoffice_output): os.unlink(expected_libreoffice_output)
        if os.path.exists(output_pdf_path) and expected_libreoffice_output != output_pdf_path: os.unlink(output_pdf_path)

        command = [
            self.libreoffice_path, '--headless', '--convert-to', 'pdf',
            '--outdir', output_dir, input_doc_path
        ]
        logging.info(f"Executando conversão: {' '.join(command)}")
        try:
            process = subprocess.run(command, capture_output=True, check=False, text=True, timeout=120)
            if process.returncode != 0:
                logging.error(f"Erro no LibreOffice (código {process.returncode}): {process.stderr}")
                return False
            if not os.path.exists(expected_libreoffice_output):
                logging.error(f"Arquivo de saída esperado '{expected_libreoffice_output}' não encontrado. Stderr: {process.stderr}")
                return False
            if expected_libreoffice_output != output_pdf_path:
                 shutil.move(expected_libreoffice_output, output_pdf_path)
                 logging.info(f"Arquivo renomeado para: {output_pdf_path}")

            if os.path.exists(output_pdf_path) and os.path.getsize(output_pdf_path) > 0:
                logging.info(f"Documento convertido para PDF: {output_pdf_path}")
                return True
            else:
                logging.error(f"Arquivo final '{output_pdf_path}' não encontrado ou vazio.")
                return False
        except subprocess.TimeoutExpired:
            logging.error("Conversão do LibreOffice excedeu o tempo limite.")
            return False
        except Exception as e:
            logging.error(f"Erro inesperado ao converter documento: {str(e)}")
            return False
        finally:
             if os.path.exists(expected_libreoffice_output) and expected_libreoffice_output != output_pdf_path:
                  try: os.unlink(expected_libreoffice_output)
                  except: pass
